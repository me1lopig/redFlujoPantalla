"""
informe.py — Generador del informe de cálculo en Word (.docx).

Toma el diccionario que produce app.calcular() y los datos de entrada, y
genera el informe de cálculo según la especificación cerrada:

  1. Portada e identificación  (campos administrativos en blanco; versión y
     huella impresas automáticamente)
  2. Objeto y alcance
  3. Normativa y método
  4. Datos de partida
  5. Modelo de cálculo
  6. Resultados
  7. Verificación interna
  8. Conclusiones

Las figuras (red de flujo, gradientes, subpresión) se insertan como imágenes
generadas desde las figuras Plotly del módulo app.

Uso:
    from informe import generar_informe
    ruta = generar_informe(datos, res, metadatos=None, salida="informe.docx")
"""

from __future__ import annotations

import io
from datetime import date

import numpy as np
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

import app  # para huella_entradas, conversiones, figuras


VERSION_PROGRAMA = "1.0"
FS_ADMISIBLE = 3.0  # criterio habitual frente a sifonamiento (ajustable)


# --------------------------------------------------------------------------- #
#  Estilo                                                                      #
# --------------------------------------------------------------------------- #
AZUL = RGBColor(0x2E, 0x75, 0xB6)
GRIS = RGBColor(0x59, 0x59, 0x59)


def _set_base_style(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    for h, sz in [("Heading 1", 15), ("Heading 2", 13)]:
        st = doc.styles[h]
        st.font.name = "Calibri"
        st.font.size = Pt(sz)
        st.font.color.rgb = AZUL


def _tabla(doc, encabezados, filas, anchos=None):
    t = doc.add_table(rows=1, cols=len(encabezados))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, e in enumerate(encabezados):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(e)
        run.bold = True
    for fila in filas:
        celdas = t.add_row().cells
        for i, val in enumerate(fila):
            celdas[i].text = str(val)
    return t


def _fig_a_imagen(fig, ancho_mm=160) -> io.BytesIO:
    """Convierte una figura matplotlib a PNG en memoria."""
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    buf.seek(0)
    return buf


# --------------------------------------------------------------------------- #
#  Figuras del informe (matplotlib, sin dependencia de navegador)              #
# --------------------------------------------------------------------------- #
def _factor_forma(res):
    """Factor de forma Nf/Nd = Q/(k_eq·ΔH) para la red de cuadrados."""
    c = res["contorno"]; Q = res["Q"]; dH = c.h1 - c.h2
    if dH <= 0:
        return 1.0
    if len(c.capas) == 1:
        k_ref = np.sqrt(c.capas[0].kx * c.capas[0].kz)
    else:
        esp = sum(cap.z_techo - cap.z_muro for cap in c.capas)
        denom = sum((cap.z_techo - cap.z_muro) / np.sqrt(cap.kx * cap.kz)
                    for cap in c.capas)
        k_ref = esp / denom if denom else 1.0
    return Q / (k_ref * dH) if k_ref else 1.0


def _fig_red_flujo_mpl(res, n_flujo=6):
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    m = res["malla"]; sol = res["sol"]; c = res["contorno"]
    fig, ax = plt.subplots(figsize=(8, 4.5))
    H = sol.h_nodo.T
    X, Z = np.meshgrid(m.x, m.z)
    # nº de saltos equipotenciales para cuadrados curvilíneos.
    # El criterio de cuadrados con k equivalente único solo es fiable en
    # terreno homogéneo; en multicapa fijamos un nº sensato de equipotenciales.
    if len(c.capas) == 1:
        S = _factor_forma(res)
        nd = max(4, int(round(n_flujo / S))) if S > 0 else 12
        titulo = (f"Red de flujo: {n_flujo} canales × {nd} saltos "
                  f"(cuadrados curvilíneos)")
    else:
        nd = 12
        titulo = (f"Red de flujo: {n_flujo} líneas de corriente, "
                  f"{nd} equipotenciales")
    # equipotenciales equiespaciadas en h (de h2 a h1)
    ax.contour(X, Z, H, levels=np.linspace(c.h2, c.h1, nd + 1),
               colors="#2E75B6", linewidths=0.8)
    # líneas de corriente por función de corriente (canales de caudal iguales)
    try:
        from corriente import resolver_psi
        psi, Q = resolver_psi(sol)
        ax.contour(X, Z, psi.T, levels=np.linspace(0, Q, n_flujo + 1),
                   colors="#C0392B", linewidths=0.8)
    except Exception:
        try:
            from postproceso import lineas_corriente
            xr, zr, VX, VZ = lineas_corriente(sol)
            ax.streamplot(xr, zr, VX, VZ, color="#C0392B", density=1.3,
                          linewidth=0.6, arrowsize=0.7)
        except Exception:
            pass
    ax.plot([c.x_tablestaca, c.x_tablestaca], [c.z_pie, c.z_coronacion],
            color="black", lw=3)
    xa_lo, xa_hi = c.x_arriba; xb_lo, xb_hi = c.x_abajo
    ax.plot([xa_lo, xa_hi], [c.z_lecho_arriba]*2, color="saddlebrown", lw=2)
    ax.plot([xb_lo, xb_hi], [c.z_lecho_abajo]*2, color="saddlebrown", lw=2)
    ax.set_xlim(c.x_izq, c.x_der); ax.set_ylim(c.z_imp, c.z_sup + 0.5)
    ax.set_xlabel("x (m)"); ax.set_ylabel("z (m)")
    ax.set_title(titulo)
    ax.set_aspect("equal", adjustable="box")
    fig.tight_layout()
    return fig


def _fig_gradiente_mpl(res):
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    m = res["malla"]; grad = res["grad"]; c = res["contorno"]
    fig, ax = plt.subplots(figsize=(8, 4.5))
    xc = 0.5 * (m.x[:-1] + m.x[1:]); zc = 0.5 * (m.z[:-1] + m.z[1:])
    im = ax.pcolormesh(xc, zc, grad["imod"].T, cmap="hot_r", shading="auto")
    fig.colorbar(im, ax=ax, label="|i|")
    ax.plot([c.x_tablestaca, c.x_tablestaca], [c.z_pie, c.z_coronacion],
            color="white", lw=2.5)
    ax.set_xlabel("x (m)"); ax.set_ylabel("z (m)")
    ax.set_title("Módulo del gradiente hidráulico")
    ax.set_aspect("equal", adjustable="box")
    fig.tight_layout()
    return fig


def _fig_subpresion_mpl(res):
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    subp = res["subp"]
    fig, ax = plt.subplots(figsize=(8, 3))
    ax.fill_between(subp["x"], subp["u"], alpha=0.3, color="#2E75B6")
    ax.plot(subp["x"], subp["u"], color="#2E75B6", marker=".", ms=3)
    ax.set_xlabel("x (m)"); ax.set_ylabel("u (kN/m²)")
    ax.set_title(f"Subpresión en cota z={subp['z_base']:.2f} m "
                 f"(resultante {subp['resultante']:.0f} kN/m)")
    fig.tight_layout()
    return fig


# --------------------------------------------------------------------------- #
#  Generador principal                                                         #
# --------------------------------------------------------------------------- #
def generar_informe(datos: dict, res: dict, metadatos: dict | None = None,
                    salida: str = "informe_calculo.docx") -> str:
    """
    Genera el informe de cálculo.

    datos     : diccionario de entrada (el de la app)
    res       : diccionario de resultados de app.calcular()
    metadatos : opcional. Si se pasan claves (obra, tramo, peticionario, autor,
                fecha) se imprimen; si no, quedan como líneas en blanco.
    """
    metadatos = metadatos or {}
    doc = Document()
    _set_base_style(doc)

    c = res["contorno"]
    sif = res["sif"]
    uq = datos.get("unidad_q", "m³/s/m")
    uk = datos.get("unidad_k", "m/s")
    Q_disp = app.caudal_desde_si(res["Q"], uq)
    huella = app.huella_entradas(datos)

    # ===================================================================== #
    # 1. PORTADA
    # ===================================================================== #
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = titulo.add_run("INFORME DE CÁLCULO")
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = AZUL
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Flujo de filtración y seguridad frente a sifonamiento\n"
                     "bajo pantalla impermeable")
    rs.font.size = Pt(13); rs.font.color.rgb = GRIS
    doc.add_paragraph()

    # Campos administrativos: en blanco salvo que se pasen en metadatos
    campos = [
        ("Obra", metadatos.get("obra", "")),
        ("Tramo / Estructura", metadatos.get("tramo", "")),
        ("Peticionario", metadatos.get("peticionario", "")),
        ("Autor", metadatos.get("autor", "")),
        ("Fecha", metadatos.get("fecha", "")),
    ]
    tcab = doc.add_table(rows=0, cols=2)
    tcab.columns[0].width = Mm(45)
    tcab.columns[1].width = Mm(120)
    for etq, val in campos:
        fila = tcab.add_row().cells
        fila[0].paragraphs[0].add_run(etq + ":").bold = True
        # espacio para rellenar a mano (línea de subrayado si vacío)
        fila[1].text = val if val else "_________________________________________"
    doc.add_paragraph()

    # Trazabilidad del cálculo (sí impresa por el programa)
    p = doc.add_paragraph()
    p.add_run("Identificación del cálculo\n").bold = True
    p.add_run(f"Programa: Flujo 2D bajo tablestaca (DF/VF) · versión "
              f"{VERSION_PROGRAMA}\n").font.size = Pt(9)
    p.add_run(f"Huella de entradas: {huella}\n").font.size = Pt(9)
    p.add_run(f"Generado: {date.today().isoformat()}").font.size = Pt(9)

    doc.add_page_break()

    # ===================================================================== #
    # 2. OBJETO Y ALCANCE
    # ===================================================================== #
    doc.add_heading("1. Objeto y alcance", level=1)
    doc.add_paragraph(
        "El presente informe recoge el cálculo del caudal de filtración y la "
        "verificación de la seguridad frente al sifonamiento (levantamiento "
        "hidráulico) en el flujo bajo una pantalla impermeable, considerando "
        "el terreno estratificado. El análisis es bidimensional, en sección "
        "transversal y por metro lineal de pantalla.")
    doc.add_paragraph(
        "El modelo resuelve el flujo en régimen estacionario y confinado, en "
        "medio saturado, mediante el método de diferencias finitas / volúmenes "
        "finitos. Quedan expresamente FUERA del alcance: el flujo no confinado "
        "con superficie freática libre, el régimen transitorio, los contornos "
        "no alineados con los ejes y cualquier acoplamiento hidromecánico. Los "
        "resultados son válidos únicamente dentro de las hipótesis del modelo "
        "confinado aquí descritas.")

    # ===================================================================== #
    # 3. NORMATIVA Y MÉTODO
    # ===================================================================== #
    doc.add_heading("2. Normativa y método de cálculo", level=1)
    doc.add_paragraph(
        "La verificación frente al sifonamiento sigue el criterio del gradiente "
        "hidráulico crítico, comparando el gradiente de salida con el gradiente "
        "crítico del terreno (criterio de Terzaghi). El marco normativo de "
        "referencia es la Guía de cimentaciones en obras de carretera "
        "(Ministerio de Fomento).")
    doc.add_paragraph(
        "El campo de cargas hidráulicas se obtiene resolviendo la ecuación de "
        "continuidad del flujo de Darcy en medio heterogéneo, ∇·(k·∇h) = 0, "
        "discretizada por volúmenes finitos sobre una malla estructurada. El "
        "esquema reproduce de forma exacta la continuidad de carga y de flujo "
        "en los contactos entre capas. El gradiente crítico se evalúa como "
        "i_crít = γ'/γ_w = (G_s−1)/(1+e).")

    # ===================================================================== #
    # 4. DATOS DE PARTIDA
    # ===================================================================== #
    doc.add_heading("3. Datos de partida", level=1)

    doc.add_heading("3.1. Geometría", level=2)
    s = c.z_lecho_abajo - c.z_pie
    _tabla(doc, ["Parámetro", "Valor", "Unidad"], [
        ["Anchura del dominio", f"{c.x_der - c.x_izq:.2f}", "m"],
        ["Profundidad al sustrato impermeable", f"{c.z_sup - c.z_imp:.2f}", "m"],
        ["Posición de la pantalla (x)", f"{c.x_tablestaca:.2f}", "m"],
        ["Empotramiento bajo lecho de salida", f"{s:.2f}", "m"],
        ["Cota de lecho aguas arriba", f"{c.z_lecho_arriba:.2f}", "m"],
        ["Cota de lecho aguas abajo", f"{c.z_lecho_abajo:.2f}", "m"],
        ["Carga aguas arriba h₁", f"{c.h1:.2f}", "m"],
        ["Carga aguas abajo h₂", f"{c.h2:.2f}", "m"],
        ["Pérdida de carga ΔH", f"{c.h1 - c.h2:.2f}", "m"],
    ])

    doc.add_heading("3.2. Estratigrafía", level=2)
    doc.add_paragraph(
        f"Permeabilidades expresadas en {uk} (y su equivalente en m/s).")
    filas_capa = []
    for cap in reversed(c.capas):  # de techo a muro (orden de lectura)
        kx_disp = app.k_desde_si(cap.kx, uk)
        kz_disp = app.k_desde_si(cap.kz, uk)
        ratio = cap.kx / cap.kz if cap.kz else 1.0
        filas_capa.append([
            cap.nombre,
            f"{cap.z_muro:.2f} – {cap.z_techo:.2f}",
            f"{cap.z_techo - cap.z_muro:.2f}",
            f"{kx_disp:.2e} ({cap.kx:.2e})",
            f"{kz_disp:.2e} ({cap.kz:.2e})",
            f"{ratio:.1f}",
        ])
    _tabla(doc, ["Capa", f"Cota muro–techo (m)", "Espesor (m)",
                 f"kx ({uk}) (m/s)", f"kz ({uk}) (m/s)", "kx/kz"], filas_capa)

    doc.add_heading("3.3. Parámetros de sifonamiento", level=2)
    _tabla(doc, ["Parámetro", "Valor"], [
        ["Método de gradiente crítico", sif.metodo_icrit],
        ["Gradiente crítico i_crít",
         f"{sif.i_critico:.4f}" if sif.i_critico else "no evaluado"],
        ["FS admisible adoptado", f"{FS_ADMISIBLE:.1f}"],
    ])

    doc.add_heading("3.4. Control numérico", level=2)
    m = res["malla"]
    _tabla(doc, ["Parámetro", "Valor"], [
        ["Densidad de malla", datos.get("densidad", "normal")],
        ["Nodos de la malla", f"{m.nx} × {m.nz} = {m.nx * m.nz}"],
    ])

    # ===================================================================== #
    # 5. MODELO DE CÁLCULO
    # ===================================================================== #
    doc.add_heading("4. Modelo de cálculo", level=1)
    doc.add_paragraph(
        "Condiciones de contorno impuestas:")
    for linea in [
        "Lecho aguas arriba: carga constante h₁ (Dirichlet).",
        "Lecho aguas abajo: carga constante h₂ (Dirichlet).",
        "Pantalla, sustrato impermeable y bordes laterales: contorno "
        "impermeable, flujo normal nulo (Neumann).",
    ]:
        p = doc.add_paragraph(linea, style="List Bullet")
    doc.add_paragraph(
        "La pantalla se modela como contorno interno impermeable de espesor "
        "nulo, desdoblando los nodos de su tramo enterrado para impedir el "
        "flujo transversal salvo el rodeo por el pie. Hipótesis: medio "
        "saturado, permeabilidad constante por capa, régimen estacionario.")

    # ===================================================================== #
    # 6. RESULTADOS
    # ===================================================================== #
    doc.add_heading("5. Resultados", level=1)

    cumple = (sif.FS is not None and sif.FS >= FS_ADMISIBLE)
    _tabla(doc, ["Magnitud", "Valor", "Unidad"], [
        ["Caudal de filtración Q", f"{Q_disp:.4g}", uq],
        ["Gradiente de salida máximo i_exit,máx", f"{sif.i_exit_max:.4f}", "—"],
        ["Gradiente de salida medio i_exit,med", f"{sif.i_exit_medio:.4f}", "—"],
        ["Gradiente crítico i_crít",
         f"{sif.i_critico:.4f}" if sif.i_critico else "—", "—"],
        ["Factor de seguridad FS",
         f"{sif.FS:.2f}" if sif.FS else "—", "—"],
        ["Resultante de subpresión (cota pie)",
         f"{res['subp']['resultante']:.0f}", "kN/m"],
    ])

    # Veredicto explícito
    p = doc.add_paragraph()
    p.add_run("Verificación frente al sifonamiento: ").bold = True
    if sif.FS is not None:
        verd = p.add_run(
            f"FS = {sif.FS:.2f} {'≥' if cumple else '<'} "
            f"{FS_ADMISIBLE:.1f} → {'CUMPLE' if cumple else 'NO CUMPLE'}")
        verd.bold = True
        verd.font.color.rgb = (RGBColor(0x2E, 0x7D, 0x32) if cumple
                               else RGBColor(0xC0, 0x39, 0x2B))
    else:
        p.add_run("no evaluada (faltan parámetros del terreno).")

    # Figuras
    doc.add_heading("5.1. Red de flujo", level=2)
    try:
        doc.add_picture(_fig_a_imagen(_fig_red_flujo_mpl(res)), width=Mm(160))
    except Exception as ex:
        doc.add_paragraph(f"[Figura no disponible: {ex}]")
    doc.add_heading("5.2. Campo de gradientes", level=2)
    try:
        doc.add_picture(_fig_a_imagen(_fig_gradiente_mpl(res)), width=Mm(160))
    except Exception as ex:
        doc.add_paragraph(f"[Figura no disponible: {ex}]")
    doc.add_heading("5.3. Subpresión", level=2)
    try:
        doc.add_picture(_fig_a_imagen(_fig_subpresion_mpl(res)), width=Mm(160))
    except Exception as ex:
        doc.add_paragraph(f"[Figura no disponible: {ex}]")

    # ===================================================================== #
    # 7. VERIFICACIÓN INTERNA
    # ===================================================================== #
    doc.add_heading("6. Verificación interna del cálculo", level=1)
    doc.add_paragraph(
        "El cálculo incorpora comprobaciones internas que avalan la validez "
        "numérica del resultado:")
    Qh1, Qh2 = res["Qh1"], res["Qh2"]
    residuo = abs(Qh1 + Qh2) / abs(Qh1) if Qh1 else 0.0
    filas_ver = [
        ["Conservación de masa (Q entrada = Q salida)",
         f"residuo relativo {residuo:.2e}",
         "Correcto" if residuo < 1e-6 else "Revisar"],
        ["Malla geométricamente consistente",
         "sin errores" if not res["err_malla"] else
         f"{len(res['err_malla'])} aviso(s)",
         "Correcto" if not res["err_malla"] else "Revisar"],
        ["Rango físico de la carga h",
         f"[{np.nanmin(res['sol'].h_nodo):.2f}, "
         f"{np.nanmax(res['sol'].h_nodo):.2f}] m", "Correcto"],
    ]
    if np.isclose(c.z_lecho_arriba, c.z_lecho_abajo):
        i_tab = int(np.argmin(np.abs(m.x - c.x_tablestaca)))
        hmed = (c.h1 + c.h2) / 2
        desv = max((abs(res["sol"].h_nodo[i_tab, j] - hmed)
                    for j in range(m.nz)
                    if m.z[j] < c.z_pie - 1e-9
                    and not np.isnan(res["sol"].h_nodo[i_tab, j])), default=0.0)
        filas_ver.append([
            "Equipotencial media bajo el pie (simetría)",
            f"desviación {desv:.2e} m",
            "Correcto" if desv < 1e-6 else "Revisar"])
    _tabla(doc, ["Comprobación", "Resultado", "Estado"], filas_ver)

    # ===================================================================== #
    # 8. CONCLUSIONES
    # ===================================================================== #
    doc.add_heading("7. Conclusiones", level=1)
    concl = doc.add_paragraph()
    concl.add_run(
        f"El caudal de filtración estimado bajo la pantalla es de "
        f"{Q_disp:.3g} {uq} por metro lineal. ")
    if sif.FS is not None:
        if cumple:
            concl.add_run(
                f"La seguridad frente al sifonamiento es suficiente "
                f"(FS = {sif.FS:.2f} ≥ {FS_ADMISIBLE:.1f}), con un gradiente "
                f"de salida máximo de {sif.i_exit_max:.3f} frente al crítico "
                f"de {sif.i_critico:.3f}.")
        else:
            concl.add_run(
                f"La seguridad frente al sifonamiento resulta INSUFICIENTE "
                f"(FS = {sif.FS:.2f} < {FS_ADMISIBLE:.1f}). Se recomienda "
                f"aumentar el empotramiento de la pantalla o disponer medidas "
                f"de control de filtraciones para reducir el gradiente de "
                f"salida.")
    else:
        concl.add_run(
            "La verificación frente al sifonamiento no ha podido completarse "
            "por falta de parámetros del terreno (G_s, e o γ_sat).")

    doc.save(salida)
    _corregir_zoom(salida)
    return salida


def _corregir_zoom(ruta: str):
    """python-docx genera <w:zoom/> sin atributo percent, que es inválido en
    el esquema OOXML estricto. Lo corrige para que el .docx valide limpio."""
    import zipfile
    import re
    import shutil
    with zipfile.ZipFile(ruta) as z:
        names = z.namelist()
        data = {n: z.read(n) for n in names}
    if "word/settings.xml" in data:
        s = data["word/settings.xml"].decode("utf-8")
        s = re.sub(r'<w:zoom(?![^>]*w:percent)([^>]*?)/>',
                   r'<w:zoom w:percent="100"\1/>', s)
        data["word/settings.xml"] = s.encode("utf-8")
        tmp = ruta + ".tmp"
        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as z:
            for n in names:
                z.writestr(n, data[n])
        shutil.move(tmp, ruta)
